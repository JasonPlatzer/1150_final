''' This program will create a random taco recipe cookbook using random recipes from an api. It will create a word document'''
import requests
import docx

#This block of code makes the front page.
document = docx.Document()   #creating a docx document
document.add_paragraph('Random Taco Cookbook', 'Title')   #adding the title of the book
document.add_picture('resizedTaco.jpg')   #adding picture
document.add_paragraph('Credits', 'Heading1')   #add the credits heading
document.add_paragraph('Photo: Miguel Andrade', 'List Bullet')   #adding first bullet point
document.add_paragraph('Tacos from random taco API URL: https://taco-1150.herokuapp.com/random/?full_taco=true', 'List Bullet')   #adding second bullet point
document.add_paragraph('Code by Jason Platzer', 'List Bullet')   #adding third bullet point
document.add_page_break()   #going to next page
#This block of code is a for loop to get the names and recipes of the components and put them into variables. Then write each chapter of the cookbook.
for responses in range(3):   #a for loop for chapters in the book
    response = requests.get('https://taco-1150.herokuapp.com/random/?full_taco=true').json()   #getting the data from api
    base = response['base_layer']['recipe']   #putting the recipe of the base layer into a variable
    seasoning = response['seasoning']['recipe']   #putting the recipe of the seasoning into a variable
    mix = response['mixin']['recipe']   #putting the recipe of the mix into a variable
    condiment = response['condiment']['recipe']   #putting the recipe of the condiment into a variable
    shell = response['shell']['recipe']   #putting the recipe of the shell into a variable
    base_name = response['base_layer']['name']   #putting name of base into a variable
    seasoning_name = response['seasoning']['name']   #putting the name of the seasoning into a variable
    mixin_name = response['mixin']['name']   #putting the name of the mix into a variable
    condiment_name = response['condiment']['name']   #putting the name of the condiment into a variable
    shell_name = response['shell']['name']   #putting the name of the shell into a variable
    p = document.add_paragraph(base_name, 'Title')# adding the name of the base to a  title paragraph in document
    p.add_run(' with ')   #adding a with to the title paragraph
    p.add_run(seasoning_name)   #adding the name of seasoning to the title paragraph
    p.add_run(' and ' )    #adding an and to the title paragraph
    p.add_run(condiment_name)   #addoing the name of the condiment to the title paragraph
    p.add_run(' in ')    #adding an in to the title paragraph
    p.add_run(shell_name)   #adding the name of the shell to the title paragraph
    document.add_paragraph('Base layer', 'Heading1')   #ddding a base layer header to go before recipe of the base layer
    document.add_paragraph(base)   #adding the recipe of the base layer to the document
    document.add_paragraph('Seasoning', 'Heading1')    #adding a seasoning header before the recipe
    document.add_paragraph(seasoning)    #adding the seasoning recipe to the document
    document.add_paragraph('Mixin', 'Heading1')     #adding a mixin header before the recipe
    document.add_paragraph(mix)    #adding the recipe of the mix to the document
    document.add_paragraph('Condiment', 'Heading1')   #adding a condiment header before the recipe
    document.add_paragraph(condiment)   #adding the recipe of the condiment to the document
    document.add_paragraph('Shell')     #adding a shell header before the recipe
    document.add_paragraph(shell)    #adding the recipe of the shell to the document
    document.add_page_break()   #adding a page break to start the next stuff on a new page
    document.save('RandomTacoRecipeBook.docx')   #saving document